from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

from config import OUTPUT_DIR


def save_summary_docx(summary_df: pd.DataFrame, filename: str = "final_summary_table.docx") -> Path:
    path = Path(OUTPUT_DIR) / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Final Summary Table", 0)
    table = doc.add_table(rows=1, cols=len(summary_df.columns))
    header_cells = table.rows[0].cells
    for i, col in enumerate(summary_df.columns):
        header_cells[i].text = str(col)
    for _, row in summary_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    doc.save(str(path))
    return path
